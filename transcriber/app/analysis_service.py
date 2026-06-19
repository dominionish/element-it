import json
import hashlib
import os
import re
import threading
import time
import uuid
from contextlib import asynccontextmanager
from datetime import datetime
from pathlib import Path
from urllib.error import HTTPError, URLError
from urllib.request import Request as UrlRequest, urlopen

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from fastapi import FastAPI, HTTPException
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from pydantic import BaseModel, Field

DATA_DIR = Path(os.getenv("DATA_DIR", "/data"))
ANALYSIS_DIR = DATA_DIR / "analysis"
INPUT_DIR = ANALYSIS_DIR / "input"
OUTPUT_DIR = ANALYSIS_DIR / "output"
JOBS_DIR = ANALYSIS_DIR / "jobs"
LOGS_DIR = DATA_DIR / "logs"

PROXYAPI_API_KEY = (
    os.getenv("PROXYAPI_API_KEY", "").strip()
    or os.getenv("OPENAI_API_KEY", "").strip()
)
PROXYAPI_CHAT_COMPLETIONS_URL = os.getenv(
    "PROXYAPI_CHAT_COMPLETIONS_URL",
    "https://api.proxyapi.ru/openai/v1/chat/completions",
).strip()
PROXYAPI_MODEL = (
    os.getenv("PROXYAPI_MODEL", "").strip()
    or os.getenv("OPENAI_MODEL", "gpt-5-mini-2025-08-07").strip()
)
PROXYAPI_TIMEOUT = int(
    os.getenv("PROXYAPI_TIMEOUT", os.getenv("OPENAI_TIMEOUT", "600"))
)
PROXYAPI_MAX_COMPLETION_TOKENS = int(
    os.getenv(
        "PROXYAPI_MAX_COMPLETION_TOKENS",
        os.getenv("OPENAI_MAX_OUTPUT_TOKENS", "12000"),
    )
)
ANALYSIS_MAX_TRANSCRIPT_CHARS = int(os.getenv("ANALYSIS_MAX_TRANSCRIPT_CHARS", "1500000"))

MEETING_ANALYSIS_PROMPT = (
    "Проанализируй стенограмму и подготовь краткое эссе и протокол совещания "
    "с таблицей с поручениями, в таблице укажи срок и ответственного за выполнение "
    "и сам текст совещания."
)

for directory in [INPUT_DIR, OUTPUT_DIR, JOBS_DIR, LOGS_DIR]:
    directory.mkdir(parents=True, exist_ok=True)


class AnalysisJobRequest(BaseModel):
    transcript: str = Field(min_length=1)
    task_id: str = ""
    project: str = ""
    source_name: str = "Стенограмма совещания"
    request_key: str = ""


@asynccontextmanager
async def lifespan(_app: FastAPI):
    analysis_log(
        "сервис анализа запущен",
        provider="proxyapi",
        endpoint=PROXYAPI_CHAT_COMPLETIONS_URL,
        model=PROXYAPI_MODEL,
        api_key_configured=bool(PROXYAPI_API_KEY),
        jobs_dir=str(JOBS_DIR),
        poll_interval_seconds=2,
    )
    threading.Thread(target=worker_loop, daemon=True).start()
    yield


app = FastAPI(title="Meeting Analysis API", version="1.0", lifespan=lifespan)


def now_iso() -> str:
    return datetime.now().isoformat(timespec="seconds")


def redact_secrets(value: str) -> str:
    value = str(value or "")
    if PROXYAPI_API_KEY:
        value = value.replace(PROXYAPI_API_KEY, "[REDACTED]")
    return re.sub(r"sk-[A-Za-z0-9_-]{8,}", "sk-[REDACTED]", value)


def analysis_log(message: str, **fields):
    safe_fields = {
        key: redact_secrets(value) if isinstance(value, str) else value
        for key, value in fields.items()
    }
    suffix = " " + json.dumps(safe_fields, ensure_ascii=False, default=str) if safe_fields else ""
    print(f"{now_iso()} [analysis] {message}{suffix}", flush=True)


def safe_name(value: str) -> str:
    value = re.sub(r'[<>:"/\\|?*]+', "_", str(value or ""))
    value = re.sub(r"\s+", " ", value).strip(" ._")
    return value or "Совещание"


def ready_file(job_id: str) -> Path:
    return JOBS_DIR / f"{job_id}.ready.json"


def status_file(job_id: str) -> Path:
    return JOBS_DIR / f"{job_id}.status.json"


def write_json(path: Path, data: dict):
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    temporary.replace(path)


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def update_status(job_key: str, **values):
    path = status_file(job_key)
    status = {}
    if path.exists():
        try:
            status = read_json(path)
        except Exception:
            status = {}
    status.update(values)
    status["updated_at"] = now_iso()
    write_json(path, status)


def response_schema() -> dict:
    return {
        "type": "object",
        "properties": {
            "essay": {"type": "string"},
            "meeting_title": {"type": "string"},
            "meeting_date": {"type": "string"},
            "participants": {"type": "array", "items": {"type": "string"}},
            "summary": {"type": "string"},
            "agenda": {"type": "array", "items": {"type": "string"}},
            "decisions": {"type": "array", "items": {"type": "string"}},
            "assignments": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "task": {"type": "string"},
                        "deadline": {"type": "string"},
                        "responsible": {"type": "string"},
                    },
                    "required": ["task", "deadline", "responsible"],
                    "additionalProperties": False,
                },
            },
        },
        "required": [
            "essay",
            "meeting_title",
            "meeting_date",
            "participants",
            "summary",
            "agenda",
            "decisions",
            "assignments",
        ],
        "additionalProperties": False,
    }


def proxyapi_request(transcript: str) -> dict:
    if not PROXYAPI_API_KEY:
        raise RuntimeError("PROXYAPI_API_KEY не настроен")

    analysis_log(
        "отправляю стенограмму в ProxyAPI Chat Completions",
        endpoint=PROXYAPI_CHAT_COMPLETIONS_URL,
        model=PROXYAPI_MODEL,
        transcript_chars=len(transcript),
        max_completion_tokens=PROXYAPI_MAX_COMPLETION_TOKENS,
        timeout_seconds=PROXYAPI_TIMEOUT,
    )

    instructions = (
        "Ты готовишь деловые документы на русском языке. Используй только факты из стенограммы. "
        "Не выдумывай участников, сроки, решения или поручения. Если срок или ответственный не названы, "
        "укажи 'Не указан'. Эссе должно быть кратким и содержательным. Протокол должен отражать ход "
        "обсуждения, решения и поручения. Верни данные строго по заданной JSON-схеме."
    )
    payload = {
        "model": PROXYAPI_MODEL,
        "messages": [
            {"role": "system", "content": instructions},
            {
                "role": "user",
                "content": f"{MEETING_ANALYSIS_PROMPT}\n\nСТЕНОГРАММА:\n{transcript}",
            },
        ],
        "max_completion_tokens": PROXYAPI_MAX_COMPLETION_TOKENS,
        "response_format": {
            "type": "json_schema",
            "json_schema": {
                "name": "meeting_analysis",
                "strict": True,
                "schema": response_schema(),
            },
        },
    }
    request = UrlRequest(
        PROXYAPI_CHAT_COMPLETIONS_URL,
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {PROXYAPI_API_KEY}",
            "Content-Type": "application/json",
            "Accept": "application/json",
            "User-Agent": "meeting-analysis-service/1.0",
        },
        method="POST",
    )
    try:
        with urlopen(request, timeout=PROXYAPI_TIMEOUT) as response:
            raw = response.read().decode("utf-8", errors="replace")
    except HTTPError as error:
        body = error.read(4000).decode("utf-8", errors="replace")
        safe_body = redact_secrets(body[:2000])
        analysis_log("ProxyAPI вернул ошибку", status_code=error.code, response=safe_body)
        raise RuntimeError(f"ProxyAPI вернул HTTP {error.code}: {safe_body}") from error
    except URLError as error:
        analysis_log("ProxyAPI недоступен", reason=str(error.reason))
        raise RuntimeError(f"ProxyAPI недоступен: {error.reason}") from error

    try:
        response_data = json.loads(raw)
    except json.JSONDecodeError as error:
        raise RuntimeError(f"ProxyAPI вернул не JSON: {raw[:500]}") from error

    choices = response_data.get("choices") or []
    if not choices:
        raise RuntimeError(f"ProxyAPI не вернул choices: {raw[:1000]}")
    choice = choices[0]
    message = choice.get("message") or {}
    if message.get("refusal"):
        raise RuntimeError(f"Модель отказалась обработать стенограмму: {message['refusal']}")
    finish_reason = choice.get("finish_reason", "")
    usage = response_data.get("usage") or {}
    analysis_log(
        "ответ ProxyAPI получен",
        response_id=response_data.get("id", ""),
        finish_reason=finish_reason,
        model=response_data.get("model", PROXYAPI_MODEL),
        input_tokens=usage.get("prompt_tokens"),
        output_tokens=usage.get("completion_tokens"),
        total_tokens=usage.get("total_tokens"),
    )

    output_text = message.get("content") or ""
    if isinstance(output_text, list):
        output_text = "".join(
            str(item.get("text") or "")
            for item in output_text
            if isinstance(item, dict)
        )
    output_text = str(output_text).strip()
    if finish_reason == "length":
        raise RuntimeError(
            "ProxyAPI остановил ответ по лимиту токенов; увеличьте PROXYAPI_MAX_COMPLETION_TOKENS"
        )
    if output_text.startswith("```"):
        output_text = re.sub(r"^```(?:json)?\s*|\s*```$", "", output_text, flags=re.IGNORECASE)
    try:
        result = json.loads(output_text)
    except json.JSONDecodeError as error:
        raise RuntimeError(f"Не удалось разобрать структурированный ответ ProxyAPI: {output_text[:1000]}") from error
    result["provider_response_id"] = response_data.get("id", "")
    result["model"] = response_data.get("model", PROXYAPI_MODEL)
    return result


def configure_document(document: Document):
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)


def add_title(document: Document, title: str):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(16)


def add_list(document: Document, title: str, values: list[str]):
    document.add_heading(title, level=1)
    if not values:
        document.add_paragraph("Не указано")
        return
    for value in values:
        document.add_paragraph(str(value), style="List Bullet")


def build_essay_document(job: dict, analysis: dict, output_path: Path):
    document = Document()
    configure_document(document)
    add_title(document, "Краткое эссе по итогам совещания")
    if job.get("project"):
        document.add_paragraph(f"Проект: {job['project']}")
    if job.get("task_id"):
        document.add_paragraph(f"Задача Planfix: {job['task_id']}")
    document.add_paragraph(analysis.get("essay") or "Содержание не сформировано")
    document.save(output_path)


def build_protocol_document(job: dict, analysis: dict, transcript: str, output_path: Path):
    document = Document()
    configure_document(document)
    add_title(document, analysis.get("meeting_title") or "Протокол совещания")
    document.add_paragraph(f"Дата: {analysis.get('meeting_date') or 'Не указана'}")
    if job.get("project"):
        document.add_paragraph(f"Проект: {job['project']}")
    if job.get("task_id"):
        document.add_paragraph(f"Задача Planfix: {job['task_id']}")

    participants = analysis.get("participants") or []
    document.add_heading("Участники", level=1)
    document.add_paragraph(", ".join(participants) if participants else "Не указаны")
    document.add_heading("Краткое содержание", level=1)
    document.add_paragraph(analysis.get("summary") or "Не указано")
    add_list(document, "Повестка и обсуждение", analysis.get("agenda") or [])
    add_list(document, "Принятые решения", analysis.get("decisions") or [])

    document.add_heading("Поручения", level=1)
    assignments = analysis.get("assignments") or []
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Поручение", "Срок", "Ответственный"]
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    if assignments:
        for assignment in assignments:
            cells = table.add_row().cells
            cells[0].text = str(assignment.get("task") or "Не указано")
            cells[1].text = str(assignment.get("deadline") or "Не указан")
            cells[2].text = str(assignment.get("responsible") or "Не указан")
    else:
        cells = table.add_row().cells
        cells[0].text = "Поручения не зафиксированы"
        cells[1].text = "Не указан"
        cells[2].text = "Не указан"

    document.add_page_break()
    document.add_heading("Стенограмма совещания", level=1)
    for block in re.split(r"\n\s*\n", transcript.strip()):
        if block.strip():
            document.add_paragraph(block.strip())
    document.save(output_path)


def process_job(job: dict):
    job_id = job["job_id"]
    try:
        transcript_path = Path(job["transcript_path"])
        transcript = transcript_path.read_text(encoding="utf-8", errors="replace").strip()
        if not transcript:
            raise RuntimeError("Стенограмма пуста")
        analysis_log(
            "начинаю анализ совещания",
            job_id=job_id,
            task_id=job.get("task_id", ""),
            source_name=job.get("source_name", ""),
            transcript_chars=len(transcript),
        )
        update_status(job_id, status="running", started_at=now_iso())

        analysis = proxyapi_request(transcript)
        output_dir = OUTPUT_DIR / job_id
        output_dir.mkdir(parents=True, exist_ok=True)
        base_name = safe_name(Path(job.get("source_name") or "Совещание").stem)[:100]
        essay_path = output_dir / f"Эссе_{base_name}.docx"
        protocol_path = output_dir / f"Протокол_{base_name}.docx"
        analysis_path = output_dir / "analysis.json"

        build_essay_document(job, analysis, essay_path)
        build_protocol_document(job, analysis, transcript, protocol_path)
        write_json(analysis_path, analysis)
        update_status(
            job_id,
            status="done",
            finished_at=now_iso(),
            essay_path=str(essay_path),
            protocol_path=str(protocol_path),
            analysis_path=str(analysis_path),
            provider_response_id=analysis.get("provider_response_id", ""),
            model=analysis.get("model", PROXYAPI_MODEL),
        )
        analysis_log(
            "анализ завершён, документы созданы",
            job_id=job_id,
            response_id=analysis.get("provider_response_id", ""),
            essay_file=essay_path.name,
            essay_size_bytes=essay_path.stat().st_size,
            protocol_file=protocol_path.name,
            protocol_size_bytes=protocol_path.stat().st_size,
        )
    except Exception as error:
        safe_error = redact_secrets(str(error))
        update_status(job_id, status="error", error=safe_error, finished_at=now_iso())
        analysis_log("анализ завершился ошибкой", job_id=job_id, error=safe_error)


def process_ready_jobs_once() -> int:
    processed = 0
    for path in sorted(JOBS_DIR.glob("*.ready.json")):
        try:
            job = read_json(path)
            consumed = JOBS_DIR / f"{job['job_id']}.processing.json"
            try:
                path.replace(consumed)
            except FileNotFoundError:
                continue
            analysis_log("задание взято из очереди", job_id=job.get("job_id", ""), queue_file=path.name)
            process_job(job)
            processed += 1
        except Exception as error:
            print(f"Ошибка чтения analysis job {path}: {error}", flush=True)
    return processed


def worker_loop():
    while True:
        try:
            process_ready_jobs_once()
        except Exception as error:
            print(f"Ошибка analysis worker: {error}", flush=True)
        time.sleep(2)


@app.get("/health")
def health():
    return {
        "status": "ok",
        "provider": "proxyapi",
        "endpoint": PROXYAPI_CHAT_COMPLETIONS_URL,
        "api_key_configured": bool(PROXYAPI_API_KEY),
        "model": PROXYAPI_MODEL,
        "data_dir": str(ANALYSIS_DIR),
    }


@app.post("/jobs")
def create_job(request: AnalysisJobRequest):
    transcript = request.transcript.strip()
    if len(transcript) > ANALYSIS_MAX_TRANSCRIPT_CHARS:
        raise HTTPException(
            status_code=413,
            detail=f"Стенограмма длиннее допустимых {ANALYSIS_MAX_TRANSCRIPT_CHARS} символов",
        )
    request_key = request.request_key.strip()
    job_id = (
        hashlib.sha256(request_key.encode("utf-8")).hexdigest()[:12]
        if request_key
        else uuid.uuid4().hex[:12]
    )
    existing_status_path = status_file(job_id)
    was_existing = existing_status_path.exists()
    if was_existing:
        existing_status = read_json(existing_status_path)
        if existing_status.get("status") in {"queued", "running", "done"}:
            analysis_log(
                "повторный запрос сопоставлен с существующим заданием",
                job_id=job_id,
                status=existing_status.get("status"),
                task_id=request.task_id,
            )
            return {
                "job_id": job_id,
                "status": existing_status.get("status"),
                "duplicate": True,
            }

    transcript_path = INPUT_DIR / f"{job_id}.txt"
    transcript_path.write_text(transcript, encoding="utf-8")
    job = {
        "job_id": job_id,
        "transcript_path": str(transcript_path),
        "task_id": request.task_id,
        "project": request.project,
        "source_name": safe_name(request.source_name),
        "request_key": request_key,
        "created_at": now_iso(),
    }
    update_status(
        job_id,
        job_id=job_id,
        status="queued",
        created_at=job["created_at"],
        task_id=request.task_id,
        source_name=job["source_name"],
        model=PROXYAPI_MODEL,
    )
    write_json(ready_file(job_id), job)
    analysis_log(
        "задание поставлено в очередь",
        job_id=job_id,
        task_id=request.task_id,
        source_name=job["source_name"],
        transcript_chars=len(transcript),
        requeued=was_existing,
    )
    return {"job_id": job_id, "status": "queued", "duplicate": False}


@app.get("/jobs/{job_id}")
def get_job(job_id: str):
    path = status_file(job_id)
    if not path.exists():
        raise HTTPException(status_code=404, detail="Задача анализа не найдена")
    return read_json(path)


@app.get("/download/{job_id}/{kind}")
def download(job_id: str, kind: str):
    if kind not in {"essay", "protocol", "analysis"}:
        raise HTTPException(status_code=400, detail="kind: essay, protocol или analysis")
    path = status_file(job_id)
    if not path.exists():
        raise HTTPException(status_code=404, detail="Задача анализа не найдена")
    status = read_json(path)
    output_value = str(status.get(f"{kind}_path") or "").strip()
    if not output_value:
        raise HTTPException(status_code=404, detail=f"Файл {kind} не найден")
    output_path = Path(output_value)
    if not output_path.is_file():
        raise HTTPException(status_code=404, detail=f"Файл {kind} не найден")
    return FileResponse(output_path, filename=output_path.name)


@app.get("/", response_class=HTMLResponse)
def index():
    return """
    <html><body style='font-family:Arial;max-width:800px;margin:40px auto'>
    <h2>Meeting Analysis API</h2>
    <p>Сервис формирует эссе и протокол совещания из готовой стенограммы.</p>
    <p>Endpoint: <code>POST /jobs</code></p>
    </body></html>
    """
