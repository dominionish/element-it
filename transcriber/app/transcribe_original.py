import argparse
import json
import math
import shutil
import subprocess
import time
from pathlib import Path

import torch
import whisper
from docx import Document
from docx.shared import Pt
from pydub import AudioSegment
from pydub.utils import make_chunks

LANGUAGE = "ru"
DEFAULT_MODEL = "medium"
CHUNK_MINUTES = 5
PAUSE_SECONDS_FOR_PARAGRAPH = 2.5
DOCX_FONT = "Times New Roman"
DOCX_SIZE = 14
OUTPUT_NAME = "Расшифровка_совещания"

MODEL_LABELS = {
    "small": "Быстрый режим",
    "medium": "Рабочий режим",
    "large-v3": "Максимальное качество",
}


def log(msg=""):
    print(msg, flush=True)


def fmt_time(sec):
    sec = int(round(sec))
    h = sec // 3600
    m = (sec % 3600) // 60
    s = sec % 60
    return f"{h:02d}:{m:02d}:{s:02d}"


def run_cmd(cmd):
    p = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True, encoding="utf-8", errors="replace")
    if p.returncode != 0:
        raise RuntimeError("Команда упала:\n" + " ".join(cmd) + "\n\n" + p.stderr)
    return p


def gpu_status():
    status = {
        "cuda_available": bool(torch.cuda.is_available()),
        "device": "cpu",
        "gpu_name": None,
        "total_vram_gb": 0.0,
        "free_vram_gb": 0.0,
    }
    if torch.cuda.is_available():
        status["device"] = "cuda"
        status["gpu_name"] = torch.cuda.get_device_name(0)
        try:
            free, total = torch.cuda.mem_get_info(0)
            status["free_vram_gb"] = round(free / (1024**3), 2)
            status["total_vram_gb"] = round(total / (1024**3), 2)
        except Exception:
            props = torch.cuda.get_device_properties(0)
            status["total_vram_gb"] = round(props.total_memory / (1024**3), 2)
    return status


def choose_safe_model(requested_model: str):
    st = gpu_status()
    model = requested_model
    reason = None

    if requested_model == "large-v3" and st["cuda_available"] and st["total_vram_gb"] and st["total_vram_gb"] < 8.0:
        model = "medium"
        reason = f"Для large-v3 мало видеопамяти: {st['total_vram_gb']} GB. Переключаю на medium."

    if requested_model == "large-v3" and not st["cuda_available"]:
        reason = "CUDA не найдена. large-v3 на CPU будет очень медленно, но запуск не запрещаю."

    return model, reason, st


def prepare_audio_ffmpeg(input_file: Path, output_wav: Path, enhance: bool):
    if shutil.which("ffmpeg") is None:
        raise RuntimeError("ffmpeg не найден внутри контейнера")

    filters = ["aresample=16000", "aformat=channel_layouts=mono"]
    if enhance:
        # Осторожная очистка речи: убрать низкий гул, крайний верх и выровнять громкость.
        filters = [
            "highpass=f=80",
            "lowpass=f=8000",
            "loudnorm=I=-16:TP=-1.5:LRA=11",
            "aresample=16000",
            "aformat=channel_layouts=mono",
        ]

    cmd = [
        "ffmpeg", "-y",
        "-i", str(input_file),
        "-vn",
        "-af", ",".join(filters),
        "-acodec", "pcm_s16le",
        str(output_wav),
    ]
    run_cmd(cmd)


def build_paragraphs(all_segments):
    paragraphs = []
    current_text = []
    current_start = None
    last_end = None

    for seg in all_segments:
        text = (seg.get("text") or "").strip()
        if not text:
            continue

        if not current_text:
            current_text = [text]
            current_start = seg["start"]
            last_end = seg["end"]
            continue

        gap = seg["start"] - last_end
        if gap >= PAUSE_SECONDS_FOR_PARAGRAPH:
            paragraphs.append({"timecode": fmt_time(current_start), "text": " ".join(current_text).strip()})
            current_text = [text]
            current_start = seg["start"]
        else:
            current_text.append(text)

        last_end = seg["end"]

    if current_text:
        paragraphs.append({"timecode": fmt_time(current_start), "text": " ".join(current_text).strip()})

    return paragraphs


def export_files(paragraphs, output_dir: Path, base_name: str, meta: dict):
    output_dir.mkdir(parents=True, exist_ok=True)
    txt_path = output_dir / f"{base_name}.txt"
    docx_path = output_dir / f"{base_name}.docx"
    meta_path = output_dir / f"{base_name}_meta.json"

    with open(txt_path, "w", encoding="utf-8") as f:
        for p in paragraphs:
            f.write(f"[{p['timecode']}] {p['text']}\n\n")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = DOCX_FONT
    style.font.size = Pt(DOCX_SIZE)
    doc.add_heading("Расшифровка совещания", level=1)

    for p in paragraphs:
        para = doc.add_paragraph()
        para.add_run(f"[{p['timecode']}] ").bold = True
        para.add_run(p["text"])

    doc.save(docx_path)
    meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")
    return txt_path, docx_path, meta_path


def transcribe(
    input_file: Path,
    output_dir: Path,
    model_name: str = DEFAULT_MODEL,
    output_name: str = OUTPUT_NAME,
    chunk_minutes: int = CHUNK_MINUTES,
    enhance_audio: bool = False,
    auto_safe_model: bool = True,
):
    started_total = time.time()
    input_file = Path(input_file)
    output_dir = Path(output_dir)

    if not input_file.exists():
        raise FileNotFoundError(f"Файл не найден: {input_file}")

    work_dir = Path("/tmp/whisper_work")
    chunks_dir = Path("/tmp/whisper_chunks")
    if work_dir.exists():
        shutil.rmtree(work_dir)
    if chunks_dir.exists():
        shutil.rmtree(chunks_dir)
    work_dir.mkdir(parents=True, exist_ok=True)
    chunks_dir.mkdir(parents=True, exist_ok=True)

    requested_model = model_name
    safe_model, model_reason, st = choose_safe_model(model_name) if auto_safe_model else (model_name, None, gpu_status())
    model_name = safe_model

    log("✅ Установка завершена. Whisper готов.")
    log("✅ Настройки приняты")
    log(f"Язык: {LANGUAGE}")
    log(f"Запрошенная модель: {requested_model} — {MODEL_LABELS.get(requested_model, '')}")
    if model_name != requested_model:
        log(f"Фактическая модель: {model_name}")
    else:
        log(f"Модель: {model_name}")
    if model_reason:
        log(f"⚠️ {model_reason}")
    log(f"Длина куска, минут: {chunk_minutes}")
    log(f"Улучшение речи ffmpeg: {'включено' if enhance_audio else 'выключено'}")
    log("")

    if st["cuda_available"]:
        log(f"✅ GPU: {st['gpu_name']}")
        log(f"✅ CUDA: OK")
        if st["total_vram_gb"]:
            log(f"✅ VRAM: всего {st['total_vram_gb']} GB, свободно {st['free_vram_gb']} GB")
    else:
        log("⚠️ CUDA не найдена. Работа пойдёт на CPU. Это медленно.")
    log("")

    log(f"✅ Файл загружен: {input_file.name}")

    wav_path = work_dir / f"{input_file.stem}_16k_mono.wav"
    log("⏳ Подготавливаю аудио через ffmpeg...")
    prepare_audio_ffmpeg(input_file, wav_path, enhance_audio)
    audio = AudioSegment.from_file(wav_path)
    duration_min = len(audio) / 60000
    log("✅ Аудио подготовлено")
    log(f"Файл: {wav_path}")
    log(f"Длительность: {duration_min:.1f} минут")
    log("")

    chunk_ms = chunk_minutes * 60 * 1000
    chunks = make_chunks(audio, chunk_ms)
    exported_chunks = []

    log("⏳ Разбиваю аудио на куски...")
    for i, chunk in enumerate(chunks):
        start_ms = i * chunk_ms
        end_ms = min((i + 1) * chunk_ms, len(audio))
        chunk_path = chunks_dir / f"chunk_{i+1:03d}_{start_ms//1000:06d}-{end_ms//1000:06d}.wav"
        chunk.export(chunk_path, format="wav")
        exported_chunks.append({"index": i + 1, "start_sec": start_ms / 1000, "end_sec": end_ms / 1000, "path": chunk_path})

    log("✅ Разбиение завершено")
    log(f"Создано кусков: {len(exported_chunks)}")
    log("")

    log(f"⏳ Загружаю модель Whisper: {model_name}")
    device = "cuda" if torch.cuda.is_available() else "cpu"
    model = whisper.load_model(model_name, device=device, download_root="/models")
    log(f"✅ Модель загружена на устройство: {device}")
    log("⏳ Начинаю расшифровку. Это самый долгий этап.")
    log("")

    all_segments = []
    total_chunks = len(exported_chunks)
    for item in exported_chunks:
        chunk_started = time.time()
        log(f"→ Обработка куска {item['index']} / {total_chunks}")
        result = model.transcribe(
            str(item["path"]),
            language=LANGUAGE,
            task="transcribe",
            verbose=False,
            fp16=torch.cuda.is_available(),
        )
        for seg in result.get("segments", []):
            all_segments.append({
                "start": float(seg["start"]) + item["start_sec"],
                "end": float(seg["end"]) + item["start_sec"],
                "text": seg["text"].strip(),
            })
        elapsed = time.time() - chunk_started
        percent = item["index"] / total_chunks * 100
        log(f"✅ Кусок {item['index']} / {total_chunks} готов за {elapsed:.1f} сек")
        log(f"Готово: {percent:.1f}%")

    all_segments.sort(key=lambda x: x["start"])
    log("")
    log("✅ Расшифровка завершена")
    log(f"Всего сегментов: {len(all_segments)}")

    paragraphs = build_paragraphs(all_segments)
    log("✅ Текст собран")
    log(f"Абзацев: {len(paragraphs)}")
    log("\nПервые 3 абзаца для проверки:\n")
    for p in paragraphs[:3]:
        log(f"[{p['timecode']}] {p['text']}\n")

    meta = {
        "input_file": input_file.name,
        "language": LANGUAGE,
        "requested_model": requested_model,
        "actual_model": model_name,
        "chunk_minutes": chunk_minutes,
        "enhance_audio": enhance_audio,
        "gpu_status": st,
        "segments": len(all_segments),
        "paragraphs": len(paragraphs),
        "duration_minutes": round(duration_min, 2),
    }
    txt_path, docx_path, meta_path = export_files(paragraphs, output_dir, output_name, meta)
    log("✅ Экспорт завершён")
    log(f"TXT: {txt_path}")
    log(f"DOCX: {docx_path}")
    log(f"META: {meta_path}")
    log(f"⏱ Общее время: {(time.time() - started_total) / 60:.1f} минут")
    return txt_path, docx_path


def main():
    parser = argparse.ArgumentParser(description="Colab original Whisper transcriber")
    parser.add_argument("input_file", help="Путь к аудио/видео внутри контейнера")
    parser.add_argument("--output-dir", default="/output")
    parser.add_argument("--model", default=DEFAULT_MODEL, choices=["small", "medium", "large-v3"])
    parser.add_argument("--output-name", default=OUTPUT_NAME)
    parser.add_argument("--chunk-minutes", type=int, default=CHUNK_MINUTES)
    parser.add_argument("--enhance-audio", action="store_true")
    parser.add_argument("--no-auto-safe-model", action="store_true")
    args = parser.parse_args()
    transcribe(
        Path(args.input_file),
        Path(args.output_dir),
        args.model,
        args.output_name,
        args.chunk_minutes,
        args.enhance_audio,
        not args.no_auto_safe_model,
    )


if __name__ == "__main__":
    main()
